"""Write domain blocks into python-docx documents."""

from __future__ import annotations

from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from docx_mcp.domain.models import DocumentBlock, ParagraphBlock, TableBlock


class ContentWriter:
    def write(
        self,
        document: DocxDocument,
        blocks: list[DocumentBlock],
        *,
        replace: bool = True,
    ) -> int:
        if replace:
            self._clear_body(document)

        count = 0
        for block in blocks:
            if isinstance(block, ParagraphBlock):
                self._write_paragraph(document, block)
                count += 1
            elif isinstance(block, TableBlock):
                self._write_table(document, block)
                count += 1
        return count

    def _clear_body(self, document: DocxDocument) -> None:
        body = document.element.body
        for child in list(body):
            if child.tag.endswith("sectPr"):
                continue
            body.remove(child)

    def _write_paragraph(
        self,
        parent: DocxDocument | _Cell,
        block: ParagraphBlock,
    ) -> Paragraph:
        paragraph = parent.add_paragraph()
        document = parent if isinstance(parent, DocxDocument) else parent.part.document

        if block.style and block.style.name:
            _apply_paragraph_style(document, paragraph, block.style.name)

        paragraph.clear()
        for run_format in block.runs:
            run = paragraph.add_run(run_format.text)
            if run_format.bold is not None:
                run.bold = run_format.bold
            if run_format.italic is not None:
                run.italic = run_format.italic
            if run_format.font_name:
                run.font.name = run_format.font_name
            if run_format.font_size_pt is not None:
                run.font.size = Pt(run_format.font_size_pt)
        return paragraph

    def _write_table(self, document: DocxDocument, block: TableBlock) -> None:
        if not block.rows:
            return

        row_count = len(block.rows)
        col_count = max(len(row) for row in block.rows)
        table = document.add_table(rows=row_count, cols=col_count)

        for row_idx, row in enumerate(block.rows):
            for col_idx, cell_block in enumerate(row):
                cell = table.rows[row_idx].cells[col_idx]
                for para_idx, para_block in enumerate(cell_block.paragraphs):
                    if para_idx == 0:
                        paragraph = cell.paragraphs[0]
                        paragraph.clear()
                        self._fill_paragraph(document, paragraph, para_block)
                    else:
                        self._write_paragraph(cell, para_block)

    def _fill_paragraph(
        self,
        document: DocxDocument,
        paragraph: Paragraph,
        block: ParagraphBlock,
    ) -> None:
        if block.style and block.style.name:
            _apply_paragraph_style(document, paragraph, block.style.name)

        for run_format in block.runs:
            run = paragraph.add_run(run_format.text)
            if run_format.bold is not None:
                run.bold = run_format.bold
            if run_format.italic is not None:
                run.italic = run_format.italic
            if run_format.font_name:
                run.font.name = run_format.font_name
            if run_format.font_size_pt is not None:
                run.font.size = Pt(run_format.font_size_pt)


def _apply_paragraph_style(
    document: DocxDocument,
    paragraph: Paragraph,
    style_name: str,
) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        _ensure_paragraph_style(document, style_name)
        try:
            paragraph.style = style_name
        except KeyError:
            pass


def _ensure_paragraph_style(document: DocxDocument, name: str) -> None:
    try:
        document.styles[name]
    except KeyError:
        document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

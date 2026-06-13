"""Shared fixtures and assertion helpers for service-layer tests."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx.document import Document as DocxDocument

from docx_mcp.adapters.content_writer import ContentWriter
from docx_mcp.adapters.docx_adapter import DocxAdapter
from docx_mcp.adapters.style_migrator import StyleMigrator
from docx_mcp.domain.models import DocumentBlock, ParagraphBlock, TableBlock
from docx_mcp.domain.style_profile import ParagraphStyleInfo, StyleProfile
from docx_mcp.services.read_service import ReadService
from docx_mcp.services.write_service import WriteService

_UNSET = object()

ASSETS_DIR = Path(__file__).parent / "assets"
PLAIN_DOCX = ASSETS_DIR / "plain.docx"
FORMAT_DOCX = ASSETS_DIR / "format.docx"


@pytest.fixture
def adapter() -> DocxAdapter:
    return DocxAdapter()


@pytest.fixture
def read_service(adapter: DocxAdapter) -> ReadService:
    return ReadService(adapter)


@pytest.fixture
def write_service(adapter: DocxAdapter) -> WriteService:
    return WriteService(adapter)


def collect_all_contents(read_service: ReadService, file_path: str, limit: int = 50) -> list[dict]:
    items: list[dict] = []
    offset = 0
    while True:
        batch = read_service.get_contents(file_path, offset=offset, limit=limit)
        items.extend(batch["items"])
        if not batch["has_more"]:
            break
        offset += limit
    return items


def collect_all_styles(read_service: ReadService, file_path: str, limit: int = 50) -> dict:
    paragraph_styles: list[dict] = []
    section: dict | None = None
    offset = 0
    while True:
        batch = read_service.get_styles(file_path, offset=offset, limit=limit)
        if offset == 0:
            section = batch.get("section")
        paragraph_styles.extend(batch["paragraph_styles"])
        if not batch["has_more"]:
            break
        offset += limit
    return {"paragraph_styles": paragraph_styles, "section": section}


def assert_blocks_equal(actual: list[DocumentBlock], expected: list[DocumentBlock]) -> None:
    assert len(actual) == len(expected), (
        f"Block count mismatch: {len(actual)} != {len(expected)}"
    )
    for idx, (actual_block, expected_block) in enumerate(zip(actual, expected)):
        assert type(actual_block) is type(expected_block), (
            f"Block {idx}: type {type(actual_block).__name__} != {type(expected_block).__name__}"
        )
        if isinstance(actual_block, ParagraphBlock) and isinstance(expected_block, ParagraphBlock):
            assert actual_block.text == expected_block.text, (
                f"Block {idx}: text mismatch"
            )
            actual_style = actual_block.style.name if actual_block.style else None
            expected_style = expected_block.style.name if expected_block.style else None
            assert actual_style == expected_style, (
                f"Block {idx}: style {actual_style!r} != {expected_style!r}"
            )
        elif isinstance(actual_block, TableBlock) and isinstance(expected_block, TableBlock):
            assert len(actual_block.rows) == len(expected_block.rows), (
                f"Block {idx}: row count mismatch"
            )
            for row_idx, (actual_row, expected_row) in enumerate(
                zip(actual_block.rows, expected_block.rows)
            ):
                assert len(actual_row) == len(expected_row), (
                    f"Block {idx} row {row_idx}: column count mismatch"
                )
                for col_idx, (actual_cell, expected_cell) in enumerate(
                    zip(actual_row, expected_row)
                ):
                    actual_texts = [p.text for p in actual_cell.paragraphs]
                    expected_texts = [p.text for p in expected_cell.paragraphs]
                    assert actual_texts == expected_texts, (
                        f"Block {idx} cell ({row_idx},{col_idx}): text mismatch"
                    )


def _style_fields_equal(actual: ParagraphStyleInfo, expected: ParagraphStyleInfo) -> bool:
    actual_dict = actual.to_dict()
    expected_dict = expected.to_dict()
    for key, expected_value in expected_dict.items():
        if key == "name":
            continue
        if expected_value is None:
            continue
        if actual_dict.get(key) != expected_value:
            return False
    return True


def assert_styles_equal(actual: StyleProfile, expected: StyleProfile) -> None:
    actual_names = actual.style_names()
    expected_names = expected.style_names()
    assert actual_names == expected_names, (
        f"Style names mismatch: {actual_names} != {expected_names}"
    )
    for name in expected_names:
        actual_style = actual.get_paragraph_style(name)
        expected_style = expected.get_paragraph_style(name)
        assert actual_style is not None, f"Missing style: {name}"
        assert expected_style is not None
        assert _style_fields_equal(actual_style, expected_style), (
            f"Style {name!r} definition mismatch: "
            f"{actual_style.to_dict()} != {expected_style.to_dict()}"
        )
    if expected.section is not None:
        assert actual.section is not None
        for field_name, expected_value in expected.section.to_dict().items():
            if expected_value is None:
                continue
            actual_value = getattr(actual.section, field_name)
            assert actual_value == pytest.approx(expected_value, rel=1e-3), (
                f"Section {field_name}: {actual_value!r} != {expected_value!r}"
            )
    else:
        assert actual.section is None


def assert_paragraph_style_field(
    profile: StyleProfile,
    name: str,
    **fields: object,
) -> None:
    style = profile.get_paragraph_style(name)
    assert style is not None, f"Style {name!r} not found"
    for field_name, expected_value in fields.items():
        actual_value = getattr(style, field_name)
        assert actual_value == expected_value, (
            f"Style {name!r}.{field_name}: {actual_value!r} != {expected_value!r}"
        )


def assert_style_field_explicit(
    profile: StyleProfile,
    name: str,
    **fields: object,
) -> None:
    """Strict field assertion — unlike assert_styles_equal, an expected None
    means the field MUST be None (explicit reset), not "skip the check"."""
    style = profile.get_paragraph_style(name)
    assert style is not None, f"Style {name!r} not found"
    for field_name, expected_value in fields.items():
        actual_value = getattr(style, field_name)
        assert actual_value == expected_value, (
            f"Style {name!r}.{field_name}: {actual_value!r} != {expected_value!r}"
        )


def get_docx_style_font_color(document: DocxDocument, name: str) -> str | None:
    """Read the resolved RGB font color of a docx paragraph style as hex, or None."""
    color = document.styles[name].font.color
    if color is None:
        return None
    try:
        rgb = color.rgb
    except Exception:
        return None
    return str(rgb) if rgb is not None else None


def assert_docx_style_field(
    document: DocxDocument,
    name: str,
    *,
    bold: object = _UNSET,
    font_color: object = _UNSET,
) -> None:
    """Assert actual OOXML state of a paragraph style definition."""
    style = document.styles[name]
    if bold is not _UNSET:
        assert style.font.bold == bold, (
            f"Style {name!r}.bold: {style.font.bold!r} != {bold!r}"
        )
    if font_color is not _UNSET:
        actual = get_docx_style_font_color(document, name)
        assert actual == font_color, (
            f"Style {name!r}.font_color: {actual!r} != {font_color!r}"
        )


def assert_paragraph_alignment(
    document: DocxDocument,
    style_name: str,
    text_contains: str,
    expected_alignment: object,
) -> None:
    """Assert the paragraph-level alignment of the first paragraph with the
    given style whose text contains text_contains."""
    for paragraph in document.paragraphs:
        style = paragraph.style
        if style is not None and style.name == style_name and text_contains in paragraph.text:
            actual = paragraph.paragraph_format.alignment
            assert actual == expected_alignment, (
                f"Paragraph {text_contains!r} ({style_name}): "
                f"alignment {actual!r} != {expected_alignment!r}"
            )
            return
    raise AssertionError(
        f"No paragraph with style {style_name!r} containing {text_contains!r}"
    )


def run_reformat_pipeline(
    adapter: DocxAdapter,
    read_service: ReadService,
) -> DocxDocument:
    """In-memory reformat: plain.docx content + format.docx styles union.

    Mirrors the agent MCP workflow (read contents -> read styles ->
    write contents -> union styles) without writing an output file to disk.
    """
    contents = collect_all_contents(read_service, str(PLAIN_DOCX))
    format_styles_dict = collect_all_styles(read_service, str(FORMAT_DOCX))

    output_doc = adapter.create_document()
    blocks = [WriteService._block_from_dict(item) for item in contents]
    ContentWriter().write(output_doc, blocks, replace=True)

    existing = adapter._style_extractor.extract(output_doc)
    incoming = StyleProfile.from_dict(format_styles_dict)
    merged = existing.union_with(incoming, master="other")
    StyleMigrator().apply(output_doc, merged)
    return output_doc
